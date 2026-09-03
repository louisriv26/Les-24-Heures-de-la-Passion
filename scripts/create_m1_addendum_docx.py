from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json, hashlib

ROOT=Path('/mnt/data/v101129_exec_strict')
html=(ROOT/'PRISTINE_V101128/index.html').read_text(encoding='utf-8')
def ex(n):
    m=f'const {n} = '; i=html.index(m)+len(m); return json.JSONDecoder().raw_decode(html[i:])[0]
TL=ex('TEXT_LIBRARY'); SPP=ex('SPEECH_PRESENTATION_PROJECTION')
items={x.get('id'):x for x in TL if isinstance(x,dict)}
def body_text(iid,n):
    it=items[iid]; body=it['body']; nums=it.get('body_stable_numbers') or list(range(1,len(body)+1)); return body[nums.index(n)]

cases=[
 {'pid':'PASSION24.TEXT.RELATED_HOUR_06.BODY.P043','n':43,'prev':42,'old':[49],'new':[],
  'reason':'Le point d’interrogation appartient à la phrase porteuse commencée au paragraphe précédent. Le break actuel à @49 sépare donc à tort le guillemet fermant du « ? » final.'},
 {'pid':'PASSION24.TEXT.RELATED_HOUR_06.BODY.P058','n':58,'prev':57,'old':[49],'new':[],
  'reason':'Même structure syntaxique : la question commence au paragraphe précédent et le « ? » extérieur clôt la phrase porteuse. Il doit rester dans le même flux visuel que la citation.'},
]
for c in cases:
    c['text']=body_text('PASSION24.TEXT.RELATED_HOUR_06',c['n'])
    c['prev_text']=body_text('PASSION24.TEXT.RELATED_HOUR_06',c['prev'])
    c['proj']=SPP[c['pid']]

BROWN=RGBColor(0x5E,0x33,0x19); BLACK=RGBColor(0x20,0x20,0x20); ACC=RGBColor(0x6D,0x4C,0x41); RED=RGBColor(0x9A,0x2F,0x2F); GREEN=RGBColor(0x2F,0x6B,0x42); GREY=RGBColor(0x66,0x66,0x66)
def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); shd=pr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); pr.append(shd)
    shd.set(qn('w:fill'),fill)
def borders(table,color='D9D9D9'):
    pr=table._tbl.tblPr; bd=pr.first_child_found_in('w:tblBorders')
    if bd is None: bd=OxmlElement('w:tblBorders'); pr.append(bd)
    for e in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single');x.set(qn('w:sz'),'7');x.set(qn('w:color'),color);bd.append(x)
def margins(cell,v=150):
    pr=cell._tc.get_or_add_tcPr(); mar=pr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar');pr.append(mar)
    for n in ('top','start','bottom','end'):
        x=OxmlElement(f'w:{n}');x.set(qn('w:w'),str(v));x.set(qn('w:type'),'dxa');mar.append(x)

def add_current_or_new(doc,title,text,proj,breaks,current=True):
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;borders(t,'E2B7B7' if current else 'B9D7C1')
    cell=t.cell(0,0);shade(cell,'FFF7F7' if current else 'F5FBF6');margins(cell,180)
    p=cell.paragraphs[0];r=p.add_run(title);r.bold=True;r.font.color.rgb=RED if current else GREEN;r.font.size=Pt(10.5)
    # Render semantic chunks + visual breaks. Only JESUS run here.
    run=proj['runs'][0]; rs,re_=int(run['start']),int(run['end'])
    points=sorted(set([0,len(text),rs,re_]+list(breaks)))
    paras=[[]]
    for a,b in zip(points[:-1],points[1:]):
        if a in breaks and paras[-1]: paras.append([])
        chunk=text[a:b]
        if not chunk: continue
        speaker='JESUS' if a>=rs and b<=re_ else 'OUTER'
        paras[-1].append((chunk,speaker))
    for chunks in paras:
        p=cell.add_paragraph();p.paragraph_format.space_after=Pt(7);p.paragraph_format.line_spacing=1.08
        for chunk,sp in chunks:
            r=p.add_run(chunk);r.font.name='Aptos';r.font.size=Pt(11);r.font.color.rgb=BROWN if sp=='JESUS' else BLACK
            if sp=='JESUS':r.bold=True;r.italic=True

D=Document();s=D.sections[0];s.top_margin=Inches(.65);s.bottom_margin=Inches(.6);s.left_margin=Inches(.75);s.right_margin=Inches(.75)
D.styles['Normal'].font.name='Aptos';D.styles['Normal'].font.size=Pt(10.5)
for nm,size in [('Title',20),('Heading 1',15),('Heading 2',12)]:
    st=D.styles[nm];st.font.name='Aptos';st.font.size=Pt(size);st.font.color.rgb=ACC
p=D.add_paragraph(style='Title');p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run('Les 24 Heures de la Passion\nAddendum de pré-validation — M1')
p=D.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Deux nouveaux cas découverts avant toute mutation du display');r.bold=True;r.font.size=Pt(12)
p=D.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Base : v101.128  |  Candidat : v101.129');r.font.color.rgb=GREY

t=D.add_table(rows=1,cols=1);borders(t,'D7C9A8');cell=t.cell(0,0);shade(cell,'FFFBEF');margins(cell,190)
p=cell.paragraphs[0];r=p.add_run('STOP M1 — aucune modification de l’application n’a été appliquée.');r.bold=True;r.font.color.rgb=RGBColor(0x7A,0x5A,0x10)
p=cell.add_paragraph('Les six ruptures déjà présentées dans le document de pré-validation initial restent inchangées et validées. L’audit exhaustif M1 a découvert exactement deux ruptures supplémentaires certaines. Conformément au verrou de validation, elles doivent être approuvées avant de figer le ledger de mutation et d’ouvrir M2.')

D.add_heading('1. Résultat exhaustif de M1',level=1)
for txt in ['107 lignes de revue / 101 groupes aliasés.','0 désaccord entre les deux lanes d’adjudication.','0 divergence d’alias.','0 candidat nouveau au second cycle de fermeture.','8 faux breaks certains au total : 6 déjà pré-validés + 2 nouveaux ci-dessous.']:
    D.add_paragraph(txt,style='List Bullet')
D.add_heading('2. Règle applicable aux deux nouveaux cas',level=1)
p=D.add_paragraph('Une fermeture de citation ne peut pas créer un paragraphe si la phrase porteuse continue encore par sa ponctuation terminale. Ici, le ');r=p.add_run('?');r.bold=True;p.add_run(' est extérieur aux guillemets et clôt la question commencée au paragraphe précédent. Le séparer du guillemet fermant par un break est donc incorrect.')

for i,c in enumerate(cases,1):
    D.add_heading(f"2.{i} — {c['pid']}",level=2)
    p=D.add_paragraph();p.add_run('Contexte précédent : ').bold=True;p.add_run(c['prev_text'])
    p=D.add_paragraph();p.add_run('Diagnostic : ').bold=True;p.add_run(c['reason'])
    add_current_or_new(D,'AUJOURD’HUI — v101.128',c['text'],c['proj'],c['old'],True)
    p=D.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('↓  correction proposée  ↓');r.bold=True;r.font.color.rgb=ACC
    add_current_or_new(D,'APRÈS MODIFICATION PROPOSÉE — v101.129',c['text'],c['proj'],c['new'],False)
    p=D.add_paragraph();p.add_run('Topologie proposée : ').bold=True;p.add_run('[49] → []  (suppression du break ; aucun replacement local, car la phrase se termine à la fin du record).')

D.add_heading('3. Mutation proposée après validation',level=1)
D.add_paragraph('Ajouter uniquement ces deux opérations au ledger déjà pré-validé :',style=None)
for c in cases:
    D.add_paragraph(f"{c['pid']} — supprimer le break @49 ; aucun break local de remplacement.",style='List Bullet')
D.add_paragraph('Le total de mutations topologiques certaines deviendrait alors : 8 faux breaks actuels corrigés (dont 3 relocations vers la vraie fin de phrase et 5 suppressions sans remplacement local). Aucun texte canonique, aucun span de locuteur et aucun offset utilisateur ne change.')

D.add_heading('4. Décision demandée',level=1)
t=D.add_table(rows=3,cols=2);borders(t,'CFCFCF')
opts=[('☐','Je valide les deux changements supplémentaires et autorise la reprise du script à M2.'),('☐','Je valide sous réserve de commentaires / modifications.'),('☐','Je ne valide pas encore ces deux changements.')]
for i,(a,b) in enumerate(opts):
    t.cell(i,0).text=a;t.cell(i,1).text=b;margins(t.cell(i,0),110);margins(t.cell(i,1),110)

for sec in D.sections:
    p=sec.footer.paragraphs[0];p.text='L24H — Addendum pré-validation M1 v101.128 → v101.129 | 3 septembre 2026';p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:r.font.size=Pt(8);r.font.color.rgb=GREY

out=Path('/mnt/data/L24H_v101129_PREVALIDATION_ADDENDUM_M1_2_NOUVEAUX_CAS_2026-09-03.docx');D.save(out)
print(out)
print(hashlib.sha256(out.read_bytes()).hexdigest())
